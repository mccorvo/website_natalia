from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


BOOK_DIR = Path(__file__).resolve().parents[1]
OUTPUT = BOOK_DIR / "The_Biodiversity_of_Mood_KDP_6x9.docx"

TITLE = "The Biodiversity of Mood"
SUBTITLE = (
    "How Modern Diets, Lost Food Diversity, and the Gut-Brain Axis "
    "May Be Changing Depression"
)
AUTHOR = "Natalia Corvo"
YEAR = "2026"

CHAPTERS = [
    (1, "Chapter_1_Why_Food_Belongs_Depression.docx", "Why Food Belongs in a Book About Depression"),
    (2, "Chapter_2_How_to_Read_Science_Food_Mood.docx", "How to Read the Science on Food and Mood"),
    (3, "Chapter_3_Depression_Is_Not_One_Disease.docx", "Depression Is Not One Disease"),
    (4, "Chapter_4_Mediterranean_Diet_Proof_of_Concept.docx", "The Mediterranean Diet as Proof of Concept"),
    (5, "Chapter_5_Ultra_processed_Foods_Sugar_Industrial_Diet.docx", "Ultra-Processed Foods, Sugar, and the Industrial Diet"),
    (6, "Chapter_6_Inflammation_Metabolism_Immuno_metabolic_Depression.docx", "Inflammation, Metabolism, and Immuno-Metabolic Depression"),
    (7, "Chapter_7_Gut_Brain_Axis_Without_Hype.docx", "The Gut-Brain Axis Without Hype"),
    (8, "Chapter_8_Nutrients_and_Supplements.docx", "Nutrients and Supplements: What Can Help and What Gets Oversold"),
    (9, "Chapter_9_From_Variety_to_Food_Biodiversity.docx", "From Variety to Food Biodiversity"),
    (10, "Chapter_10_The_Biology_of_a_Narrow_Diet.docx", "The Biology of a Narrow Diet"),
]

CAPTION_RE = re.compile(r"^(Box|Figure|Table)\s+\d+\.\d+\s+-\s+.+")
DRAFT_RE = re.compile(r"Working manuscript draft", re.I)


def iter_blocks(doc: Document) -> Iterable[Paragraph | Table]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def raw_clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def clean_text(text: str) -> str:
    text = raw_clean_text(text)
    if "operating brief" in text.lower():
        return ""
    replacements = {
        "This is the line that will carry into Chapter 11:": "This is the line that carries the argument forward:",
        "A biology chapter therefore leads directly into Chapter 11.": "A biology chapter therefore leads directly into the practical question of how depression narrows the menu.",
        "This is why Chapter 11 will move from the gut to the depressed menu.": "This is why the argument must move from the gut to the depressed menu.",
        "The next chapter moves from biology to behaviour.": "The argument now moves from biology to behaviour.",
        "Chapter 11 asks how the menu becomes narrow in the first place.": "The behavioural question asks how the menu becomes narrow in the first place.",
        "Chapter 11 asks how the menu becomes narrow in the first place": "the behavioural question asks how the menu becomes narrow in the first place",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("Chapter 11", "the next stage of the argument")
    return text


def style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def is_heading(paragraph: Paragraph) -> bool:
    name = style_name(paragraph).lower()
    text = clean_text(paragraph.text)
    if CAPTION_RE.match(text):
        return False
    return "heading" in name or name in {"title", "subtitle"}


def is_list(paragraph: Paragraph) -> bool:
    return "list" in style_name(paragraph).lower()


def has_real_text(paragraph: Paragraph) -> bool:
    return bool(clean_text(paragraph.text))


def source_blocks_after_draft_header(path: Path) -> list[Paragraph | Table]:
    doc = Document(path)
    blocks = list(iter_blocks(doc))
    start = 0
    for index, block in enumerate(blocks):
        if isinstance(block, Paragraph) and DRAFT_RE.search(block.text):
            start = index + 1
            break
    if start == 0:
        seen = 0
        for index, block in enumerate(blocks):
            if isinstance(block, Paragraph) and has_real_text(block):
                seen += 1
            if seen >= 4:
                start = index + 1
                break
    return blocks[start:]


def set_cell_shading(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell: _Cell, color: str = "8BA28F", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_box(paragraph: Paragraph, fill: str, border: str, top: bool = False, bottom: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    wanted = ["left", "right"]
    if top:
        wanted.append("top")
    if bottom:
        wanted.append("bottom")
    for edge in wanted:
        tag = "w:" + edge
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8" if edge == "left" else "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), border)


def set_cell_margins(cell: _Cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = tc_pr.first_child_found_in("w:tcMar")
    if margin is None:
        margin = OxmlElement("w:tcMar")
        tc_pr.append(margin)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margin.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            margin.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table: Table, width_inches: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_table_grid(table: Table, widths_inches: list[float]) -> None:
    tbl = table._tbl
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    tbl.insert(1, grid)


def set_cell_width(cell: _Cell, width_inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph: Paragraph, field_code: str, display: str = "") -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), field_code)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn("w:" + attr), "Garamond")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(r_fonts)
    r_pr.append(size)
    run.append(r_pr)
    text = OxmlElement("w:t")
    text.text = display or "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def add_page_number(paragraph: Paragraph) -> None:
    add_field(paragraph, "PAGE", "1")


def set_rfonts(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn("w:" + attr), font_name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)

    settings = doc.settings.element
    if settings.find(qn("w:mirrorMargins")) is None:
        settings.append(OxmlElement("w:mirrorMargins"))

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Garamond"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Inches(0.2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.04

    def ensure_paragraph_style(name: str, base: str = "Normal"):
        if name in styles:
            return styles[name]
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base]
        return style

    body_first = ensure_paragraph_style("Book First Paragraph")
    body_first.font.name = "Garamond"
    body_first.font.size = Pt(10.5)
    body_first.paragraph_format.first_line_indent = Inches(0)
    body_first.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_first.paragraph_format.space_after = Pt(0)
    body_first.paragraph_format.line_spacing = 1.04

    heading_1 = styles["Heading 1"]
    heading_1.font.name = "Garamond"
    heading_1.font.size = Pt(18)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(35, 52, 43)
    heading_1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_1.paragraph_format.space_before = Pt(0)
    heading_1.paragraph_format.space_after = Pt(20)
    heading_1.paragraph_format.keep_with_next = True
    heading_1.paragraph_format.first_line_indent = Inches(0)

    heading_2 = styles["Heading 2"]
    heading_2.font.name = "Garamond"
    heading_2.font.size = Pt(13.5)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(35, 52, 43)
    heading_2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_2.paragraph_format.space_before = Pt(18)
    heading_2.paragraph_format.space_after = Pt(6)
    heading_2.paragraph_format.keep_with_next = True
    heading_2.paragraph_format.first_line_indent = Inches(0)

    heading_3 = styles["Heading 3"]
    heading_3.font.name = "Garamond"
    heading_3.font.size = Pt(11.5)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(35, 52, 43)
    heading_3.paragraph_format.space_before = Pt(12)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True
    heading_3.paragraph_format.first_line_indent = Inches(0)

    for name, size, bold, italic in [
        ("Book Half Title", 20, False, False),
        ("Book Title Page Title", 28, True, False),
        ("Book Subtitle", 13, False, True),
        ("Book Author", 14, False, False),
        ("Book Front Heading", 16, True, False),
        ("Book Chapter Number", 10, False, False),
        ("Book Caption", 9.5, False, True),
        ("Book Box Title", 9.5, True, False),
        ("Book Box Body", 9.2, False, False),
        ("Book TOC Entry", 10.5, False, False),
        ("Book Header Footer", 8.5, False, False),
        ("Book Copyright", 9.2, False, False),
    ]:
        style = ensure_paragraph_style(name)
        style.font.name = "Garamond"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        style.paragraph_format.first_line_indent = Inches(0)
        style.paragraph_format.line_spacing = 1.0
        if name in {"Book Half Title", "Book Title Page Title", "Book Subtitle", "Book Author", "Book Chapter Number"}:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name == "Book Caption":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            style.paragraph_format.space_before = Pt(10)
            style.paragraph_format.space_after = Pt(4)
        if name == "Book Box Body":
            style.paragraph_format.space_after = Pt(3)


def configure_section(section) -> None:
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def unlink_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False


def setup_body_section(section) -> None:
    configure_section(section)
    section.different_first_page_header_footer = True
    unlink_header_footer(section)

    header = section.header.paragraphs[0]
    header.style = "Book Header Footer"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.text = "{} | {}".format(TITLE, AUTHOR)

    footer = section.footer.paragraphs[0]
    footer.style = "Book Header Footer"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")


def setup_front_section(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section)
    section.different_first_page_header_footer = True
    unlink_header_footer(section)
    for part in (
        section.header,
        section.footer,
        section.first_page_header,
        section.first_page_footer,
    ):
        part.paragraphs[0].text = ""


def add_centered_spacer(doc: Document, count: int = 4) -> None:
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)


def add_front_matter(doc: Document, toc_pages: dict[str, int] | None) -> None:
    setup_front_section(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.author = AUTHOR
    doc.core_properties.subject = "Food biodiversity, diet, and depression"
    doc.core_properties.comments = "KDP 6 x 9 paperback manuscript interior."

    add_centered_spacer(doc, 10)
    p = doc.add_paragraph(TITLE, style="Book Half Title")
    p.paragraph_format.space_after = Pt(0)

    doc.add_page_break()
    add_centered_spacer(doc, 7)
    p = doc.add_paragraph(TITLE, style="Book Title Page Title")
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph(SUBTITLE, style="Book Subtitle")
    p.paragraph_format.space_after = Pt(40)
    p = doc.add_paragraph("by", style="Book Author")
    p.paragraph_format.space_after = Pt(8)
    doc.add_paragraph(AUTHOR, style="Book Author")

    doc.add_page_break()
    add_centered_spacer(doc, 5)
    copyright_lines = [
        "Copyright (c) {} {}".format(YEAR, AUTHOR),
        "All rights reserved.",
        "",
        "ISBN-13: 978-X-XXXX-XXXX-X",
        "",
        "This manuscript is prepared as a paperback interior file for review and publication setup.",
    ]
    for line in copyright_lines:
        p = doc.add_paragraph(line, style="Book Copyright")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()
    p = doc.add_paragraph("A Note to the Reader", style="Book Front Heading")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_paragraphs = [
        "This book is for educational and informational purposes only. It does not provide medical diagnosis, treatment, psychotherapy, nutrition therapy, or individual clinical advice.",
        "Depression can be serious and sometimes life-threatening. Anyone experiencing suicidal thoughts, self-harm urges, psychosis, mania, severe functional decline, or rapid worsening should seek urgent professional help or contact local emergency services.",
        "Diet is treated here as one part of a larger terrain: clinical care, sleep, medications, trauma, poverty, social connection, food access, physical illness, and the food system all matter. No dietary pattern, food, nutrient, supplement, or microbiome strategy is presented as a standalone cure for depression.",
    ]
    for text in note_paragraphs:
        p = doc.add_paragraph(text, style="Book First Paragraph")
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()
    p = doc.add_paragraph("Contents", style="Book Front Heading")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    for number, _, title in CHAPTERS:
        p = doc.add_paragraph(style="Book TOC Entry")
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.right_indent = Inches(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(4.65), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        p.add_run("Chapter {}: ".format(number)).bold = True
        p.add_run(title)
        page_text = str(toc_pages.get(str(number), "")) if toc_pages else ""
        p.add_run("\t" + page_text)


def clone_run_format(src_run, dst_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.superscript = src_run.font.superscript
    dst_run.font.subscript = src_run.font.subscript
    set_rfonts(dst_run, "Garamond")
    if src_run.font.size:
        dst_run.font.size = src_run.font.size


def copy_runs(src: Paragraph, dst: Paragraph, fallback_text: str | None = None) -> None:
    raw_text = raw_clean_text(src.text)
    cleaned_text = clean_text(src.text)
    if cleaned_text != raw_text:
        if cleaned_text:
            dst.add_run(cleaned_text)
        return
    copied = False
    for src_run in src.runs:
        if not src_run.text:
            continue
        text = src_run.text.replace("\u00a0", " ")
        if not text:
            continue
        dst_run = dst.add_run(text)
        clone_run_format(src_run, dst_run)
        copied = True
    if not copied:
        dst.add_run(fallback_text if fallback_text is not None else clean_text(src.text))


def add_body_paragraph(doc: Document, src: Paragraph, first_after_break: bool) -> bool:
    text = clean_text(src.text)
    if not text:
        return first_after_break
    if is_heading(src):
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.first_line_indent = Inches(0)
        p.add_run(text)
        return True
    if is_list(src):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(2)
        copy_runs(src, p, text)
        return False
    p = doc.add_paragraph(style="Book First Paragraph" if first_after_break else "Normal")
    copy_runs(src, p, text)
    return False


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="Book Caption")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.keep_with_next = True


def add_box(doc: Document, title: str, paragraphs: list[Paragraph], kind: str) -> None:
    fill = "F1F5F0" if kind == "Box" else "F7F4EA"
    border = "7D987F" if kind == "Box" else "B39B68"
    before = doc.add_paragraph()
    before.paragraph_format.space_after = Pt(2)
    created: list[Paragraph] = []
    first = doc.add_paragraph(style="Book Box Title")
    first.paragraph_format.keep_with_next = True
    first.paragraph_format.first_line_indent = Inches(0)
    first.paragraph_format.left_indent = Inches(0.16)
    first.paragraph_format.right_indent = Inches(0.12)
    first.paragraph_format.space_before = Pt(0)
    first.paragraph_format.space_after = Pt(2)
    first.add_run(title)
    set_paragraph_box(first, fill, border, top=True, bottom=False)
    created.append(first)
    for src in paragraphs:
        text = clean_text(src.text)
        if not text:
            continue
        p = doc.add_paragraph(style="Book Box Body")
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.right_indent = Inches(0.12)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        if is_list(src):
            p.style = "List Bullet"
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.first_line_indent = Inches(-0.12)
        copy_runs(src, p, text)
        set_paragraph_box(p, fill, border, top=False, bottom=False)
        created.append(p)
    if created:
        set_paragraph_box(created[-1], fill, border, top=False, bottom=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def maybe_table_box(doc: Document, table: Table) -> bool:
    if len(table.rows) != 1 or len(table.columns) != 1:
        return False
    text = table.cell(0, 0).text.strip()
    if not text.startswith("Box "):
        return False
    lines = [clean_text(line) for line in text.splitlines() if clean_text(line)]
    if not lines:
        return False

    class SimpleParagraph:
        def __init__(self, text: str):
            self.text = text
            self.runs = []
            self.style = None

    add_box(doc, lines[0], [SimpleParagraph(line) for line in lines[1:]], "Box")
    return True


def copy_table(doc: Document, src_table: Table) -> None:
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    if rows == 0 or cols == 0:
        return
    header = [clean_text(cell.text) for cell in src_table.rows[0].cells]
    data_rows = src_table.rows[1:] if rows > 1 else src_table.rows
    before = doc.add_paragraph()
    before.paragraph_format.space_after = Pt(2)
    for r_idx, src_row in enumerate(data_rows):
        row_paragraphs: list[Paragraph] = []
        for c_idx, src_cell in enumerate(src_row.cells):
            value = clean_text(src_cell.text)
            if not value:
                continue
            p = doc.add_paragraph(style="Book Box Body")
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            label = header[c_idx] if c_idx < len(header) and header[c_idx] else "Column {}".format(c_idx + 1)
            run = p.add_run(label + ": ")
            run.bold = True
            set_rfonts(run, "Garamond")
            p.add_run(value)
            for run in p.runs:
                set_rfonts(run, "Garamond")
                run.font.size = Pt(8.8)
            row_paragraphs.append(p)
        for idx, p in enumerate(row_paragraphs):
            set_paragraph_box(
                p,
                "F7F8F4" if r_idx % 2 == 0 else "FFFFFF",
                "ADB8A8",
                top=idx == 0,
                bottom=idx == len(row_paragraphs) - 1,
            )
        if row_paragraphs:
            gap = doc.add_paragraph()
            gap.paragraph_format.space_after = Pt(2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(5)


def should_group_body(block: Paragraph | Table) -> bool:
    if isinstance(block, Table):
        return False
    text = clean_text(block.text)
    if not text:
        return True
    if CAPTION_RE.match(text) or is_heading(block):
        return False
    return True


def add_chapter(doc: Document, number: int, title: str, source_path: Path, page_break: bool) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph("CHAPTER {}".format(number), style="Book Chapter Number")
    p.paragraph_format.space_before = Pt(42)
    p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph(title, style="Heading 1")
    p.paragraph_format.space_after = Pt(26)

    blocks = source_blocks_after_draft_header(source_path)
    first_after_break = True
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if isinstance(block, Table):
            if not maybe_table_box(doc, block):
                copy_table(doc, block)
            first_after_break = True
            index += 1
            continue

        text = clean_text(block.text)
        if not text:
            index += 1
            continue
        if DRAFT_RE.search(text) or text == TITLE or text == "Chapter {}".format(number):
            index += 1
            continue

        caption = CAPTION_RE.match(text)
        if caption:
            kind = caption.group(1)
            if kind == "Table" and index + 1 < len(blocks) and isinstance(blocks[index + 1], Table):
                add_caption(doc, text)
                copy_table(doc, blocks[index + 1])
                first_after_break = True
                index += 2
                continue
            if kind in {"Box", "Figure"}:
                grouped: list[Paragraph] = []
                j = index + 1
                while j < len(blocks) and should_group_body(blocks[j]):
                    if isinstance(blocks[j], Paragraph) and has_real_text(blocks[j]):
                        grouped.append(blocks[j])
                    j += 1
                add_box(doc, text, grouped, kind)
                first_after_break = True
                index = j
                continue
            add_caption(doc, text)
            first_after_break = True
            index += 1
            continue

        first_after_break = add_body_paragraph(doc, block, first_after_break)
        index += 1


def build(output: Path, toc_pages: dict[str, int] | None = None) -> None:
    doc = Document()
    configure_document(doc)
    add_front_matter(doc, toc_pages)
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    setup_body_section(body_section)
    for number, filename, title in CHAPTERS:
        add_chapter(doc, number, title, BOOK_DIR / filename, page_break=number != 1)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", default=str(OUTPUT))
    parser.add_argument("--toc-pages-json", default="")
    args = parser.parse_args()
    toc_pages = None
    if args.toc_pages_json:
        toc_pages = json.loads(Path(args.toc_pages_json).read_text())
    build(Path(args.out), toc_pages)
    print(args.out)


if __name__ == "__main__":
    main()
